# scripts/build_qa_doc.py
# Generates the Credentia Q&A defense document (.docx) with cost + timeline tables.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT   = RGBColor(0x4F, 0x46, 0xE5)
DARK     = RGBColor(0x0F, 0x17, 0x2A)
SECOND   = RGBColor(0x47, 0x55, 0x69)
MUTED    = RGBColor(0x94, 0xA3, 0xB8)
AMBER    = RGBColor(0x92, 0x6A, 0x0A)
HEADER_BG = "0F1117"
ALT_BG    = "EEF2FF"
FONT = "Calibri"

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = SECOND

# narrow-ish margins
for sec in doc.sections:
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.7);  sec.bottom_margin = Inches(0.7)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def run(p, text, size=11, color=SECOND, bold=False, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = FONT
    return r


def heading(text, size=15, color=ACCENT, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run(p, text, size=size, color=color, bold=True)
    return p


def body(text, size=11, color=SECOND, bold=False, italic=False, space_after=6, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    run(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def labeled(label, text, label_color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, label + "  ", size=11, color=label_color, bold=True)
    run(p, text, size=11, color=SECOND)
    return p

# ── Title ─────────────────────────────────────────────────────
t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(2)
run(t, "Credentia — Q&A Defense Briefing", size=24, color=DARK, bold=True)
st = doc.add_paragraph(); st.paragraph_format.space_after = Pt(2)
run(st, "Demo Block (Slides 8–10) & Roadmap (Slides 13–15)", size=13, color=ACCENT, bold=True)
fr = doc.add_paragraph(); fr.paragraph_format.space_after = Pt(10)
run(fr, "With token-usage cost model and roadmap effort estimates.  Confidential.", size=10, color=MUTED, italic=True)

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(12)
run(note, "Presenter framing:  ", size=11, color=AMBER, bold=True)
run(note, "Every dollar and timeline figure below is a planning estimate built on today's "
          "model pricing and reasonable usage assumptions. Say the word “estimate” out loud "
          "— it builds credibility, not doubt.", size=11, color=SECOND, italic=True)

# ── Q&A items ─────────────────────────────────────────────────
qa = [
    ("Q1. “How do I know your AI isn't fabricating things about my employees?”",
     "Because the AI doesn't generate the numbers — our code does. Completion rates, revenue "
     "impact, sentiment trends, and retention signals are all computed deterministically from "
     "verified data. The model's only job is clustering free-text and writing narrative. It is "
     "architecturally unable to invent a metric or flag a named person.",
     "We even overwrite the model's copy of the figures with our ground-truth numbers after it "
     "responds — so a hallucinated number can't survive to the UI. Retention flags are pure math "
     "on a sentiment slope; the AI never sees them as something to invent."),

    ("Q2. “Won't AI costs explode as you scale to thousands of employees?”",
     "No — and this surprises people. Our AI calls are small and infrequent. A full advisory pass "
     "on one employee is about 4,000 tokens, roughly three cents. A leadership report is about "
     "three and a half cents. Even a 1,000-person org running quarterly is tens of dollars a year "
     "— a rounding error against enterprise seat pricing.",
     "Three things keep it there. One: most operations run on a cadence — quarterly, weekly — not "
     "per-keystroke. Two: prompt caching cuts repeated system-prompt input cost by ~90%. Three: we "
     "route by job — Haiku for high-volume monitoring, Sonnet for synthesis, Opus only for heavy "
     "planning. See Table 1: even with the Phase 2 features on, we model under ~$500/org/year in tokens."),

    ("Q3. “What's actually real today versus a mockup?”",
     "The core is live: the verification engine, multi-tenant isolation, the audit trail, every "
     "dashboard, provisioning, the swipe-to-verify flow you just saw. The AI layer is wired and "
     "running on Claude. Two things are deliberately mocked — the payment processor and the "
     "dedicated-database isolation mode — both scaffolded in schema, neither on the critical path.",
     "Nothing you saw in the demo was faked. The mocks are commercial plumbing — billing and "
     "physical DB separation — that we'll harden for GA, not core product risk."),

    ("Q4. “What stops Workday, Rippling, or LinkedIn from just copying this?”",
     "The separation of verified fact from AI estimate isn't a feature you bolt on — it's a "
     "schema-level commitment that touches every table and every screen. Incumbents have years of "
     "blended data they can't retroactively un-mix. And our value compounds with a network: every "
     "company that joins makes a portable credential worth more. That's a moat that grows, not one "
     "they can ship in a sprint.",
     "Their data model is their liability here. Re-architecting a live HRIS to separate attested "
     "fact from inference, retroactively, is a multi-year migration with no clean answer for "
     "historical data."),

    ("Q5. “Scoring employees with AI — isn't that a legal and ethical minefield?”",
     "We designed for exactly this fear. The AI never decides — no promotion, no comp change, no "
     "termination, no rating. It outputs ranges and suggestions, clearly labeled amber, with a "
     "human-decides disclaimer enforced in the prompt and the architecture. Every action is on a "
     "tamper-evident audit trail. That makes us more defensible in an audit, not less.",
     "Our Phase 3 pay-equity audit engine actually turns this into a compliance asset — continuous, "
     "regulator-ready fairness monitoring on top of attested data."),

    ("Q6. “How long until the roadmap features I care about are real?”",
     "Phase 2 is a 2-to-3 quarter horizon, sequenced by leverage. The quick, high-trust wins — "
     "explainable AI panels, natural-language workforce queries — are 3-to-8 week builds because "
     "they sit on data we already capture. The heavier network features, like cross-company "
     "credential portability, are 8-to-12 weeks because the hard part is the consent and legal "
     "model, not the code.",
     "Phase 3 — the talent graph, the marketplace, predictive planning — is a 12-to-24 month arc, "
     "and several pieces only unlock once we have enough tenants to make benchmarking and the "
     "network meaningful. See Table 2 for the sequencing."),

    ("Q7. “Your AI runs on a third party. What's your exposure if pricing or availability changes?”",
     "Our prompts are model-portable — structured JSON in, structured JSON out — so we're not "
     "locked to one vendor. And because our spend is so low per org, even a significant price move "
     "is immaterial to unit economics. We can also down-route non-critical work to cheaper models "
     "without touching the product.",
     "The deterministic metrics — the part that matters most for trust — don't depend on any LLM "
     "at all. The model is the narrator, not the source of truth. If it went away tomorrow, the "
     "verified platform still stands."),

    ("Q8. “What's the single biggest technical risk in the roadmap?”",
     "Honestly? Cross-company credential portability — Phase 2's marquee feature. The engineering "
     "is tractable; the consent, privacy, and cross-tenant trust model is the real work. We've "
     "budgeted 8-to-12 weeks and we're treating the legal design as a first-class part of it.",
     "Everything downstream — the talent graph, the marketplace — depends on getting that consent "
     "model right once. So we're front-loading it and over-investing in it deliberately."),
]

heading("The Toughest Questions", size=17, color=DARK, space_before=6)
for q, spoken, pressed in qa:
    heading(q, size=12.5, color=ACCENT, space_before=12, space_after=4)
    labeled("Spoken (30s):", spoken)
    labeled("If pressed:", pressed, label_color=AMBER)

# ── Table 1 — Token & Cost Model ──────────────────────────────
doc.add_page_break()
heading("Table 1 — AI Token & Cost Model", size=16, color=DARK, space_before=4)
body("Assumes current Claude Sonnet pricing (~$3 / 1M input, ~$15 / 1M output); high-volume "
     "monitoring routed to Haiku. Figures are per-org estimates.", size=10, color=MUTED, italic=True, space_after=8)

t1_head = ["Operation", "Cadence", "~Tokens (in / out)", "~Cost / run", "~Cost / 1,000-seat org / yr"]
t1_rows = [
    ("TODAY", "", "", "", ""),
    ("Employee advisory (full)", "Quarterly / employee", "2.5K / 1.5K", "~$0.03", "~$120"),
    ("Leadership report", "Weekly / team", "4K / 1.5K", "~$0.035", "~$90"),
    ("PHASE 2 ADDITIONS", "", "", "", ""),
    ("NL workforce query (RAG)", "On demand (~20/day)", "5K / 0.8K", "~$0.03", "~$220"),
    ("Explainable AI panel", "Reuses existing call", "+0.5K", "~$0 marginal", "negligible"),
    ("Compliance agent (Haiku)", "Hourly audit scan", "2K / 0.3K", "~$0.001", "~$60"),
    ("PHASE 3 ADDITIONS", "", "", "", ""),
    ("Predictive planning (Opus)", "Monthly / org", "10K / 4K", "~$0.09", "~$1"),
    ("ESTIMATED ALL-IN, FULLY ENABLED", "", "", "", "≈ $450–550 / org / yr"),
]

tbl = doc.add_table(rows=1, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(t1_head):
    shade(hdr[i], HEADER_BG)
    p = hdr[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    run(p, h, size=9.5, color=RGBColor(0xFF,0xFF,0xFF), bold=True)

section_rows = {"TODAY", "PHASE 2 ADDITIONS", "PHASE 3 ADDITIONS"}
for r in t1_rows:
    cells = tbl.add_row().cells
    is_section = r[0] in section_rows
    is_total = r[0].startswith("ESTIMATED")
    for i, val in enumerate(r):
        if is_section: shade(cells[i], "1E2433")
        elif is_total: shade(cells[i], ALT_BG)
        p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        col = RGBColor(0xC7,0xD2,0xFE) if is_section else (ACCENT if is_total else SECOND)
        run(p, val, size=9.5, color=col, bold=(is_section or is_total or i==0))

body("", space_after=2)
hl = doc.add_paragraph(); hl.paragraph_format.space_before = Pt(8)
run(hl, "Headline:  ", size=11, color=DARK, bold=True)
run(hl, "Even with every AI feature on, we model under ~$500 in tokens per customer per year. "
        "Token cost is not our scaling risk.", size=11, color=SECOND, italic=True)
body("Two cost levers in reserve: prompt caching (~90% off repeated input) and model routing "
     "(Haiku / Sonnet / Opus by job).", size=11, color=SECOND, space_after=4)

# ── Table 2 — Roadmap Effort & Timeline ───────────────────────
doc.add_page_break()
heading("Table 2 — Roadmap Effort & Timeline", size=16, color=DARK, space_before=4)
body("Engineering estimates assuming current team; the gating factor is noted where it isn't code.",
     size=10, color=MUTED, italic=True, space_after=8)

t2_head = ["Feature", "Phase", "Est. build", "Gating factor"]
t2_rows = [
    ("Explainable AI source panels", "2", "3–4 wks", "Surfacing data we already store"),
    ("NL workforce query", "2", "6–8 wks", "RAG guardrails over verified tables"),
    ("Mobile app (pulse + approvals)", "2", "8–10 wks", "Native build + review"),
    ("Integrations hub (Slack/Teams/Workday)", "2", "6–10 wks", "Per-connector, ongoing"),
    ("Third-party verifier marketplace", "2", "10–12 wks", "Partner onboarding model"),
    ("Cross-company credential portability", "2", "8–12 wks", "Consent / legal model (TOP RISK)"),
    ("Cryptographic credential anchoring", "3", "6–8 wks", "Standard must be defined first"),
    ("Pay-equity audit engine", "3", "8–10 wks", "Builds on existing equity score"),
    ("Predictive workforce planning", "3", "~1 quarter", "Needs longitudinal data"),
    ("Benchmarking-as-a-service", "3", "~1 quarter", "Needs tenant critical mass"),
    ("Developer / verification API", "3", "~1 quarter", "Versioning + rate limiting"),
    ("Autonomous compliance agent", "3", "1–2 quarters", "Builds on audit trail"),
]

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(t2_head):
    shade(hdr2[i], HEADER_BG)
    p = hdr2[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    run(p, h, size=10, color=RGBColor(0xFF,0xFF,0xFF), bold=True)

for ri, r in enumerate(t2_rows):
    cells = tbl2.add_row().cells
    is_risk = "TOP RISK" in r[3]
    for i, val in enumerate(r):
        if ri % 2 == 1: shade(cells[i], "F8FAFC")
        if is_risk: shade(cells[i], "FFFBEB")
        p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        col = AMBER if is_risk else (ACCENT if i==1 else SECOND)
        run(p, val, size=10, color=col, bold=(i==0 or is_risk))

body("", space_after=2)
body("Phase 2 total horizon: ~2–3 quarters (features parallelize across the team).",
     size=11, color=SECOND, bold=True, space_after=3)
body("Phase 3 total horizon: ~12–24 months; several items unlock only at tenant scale.",
     size=11, color=SECOND, bold=True, space_after=10)

# ── One line to memorize ──────────────────────────────────────
box = doc.add_paragraph(); box.paragraph_format.space_before = Pt(10)
run(box, "The One Line to Memorize", size=13, color=DARK, bold=True)
ol = doc.add_paragraph()
ol.paragraph_format.left_indent = Inches(0.3)
run(ol, "“The AI is cheap because it narrates instead of computes — under ~$500 per org per "
        "year fully loaded. The roadmap is sequenced by leverage — quick trust wins first, the "
        "network play next — and our biggest risk is a consent model, not a line of code.”",
    size=12, color=ACCENT, italic=True, bold=True)

out = r"C:\Users\tyrel\credentia\Credentia_QA_Defense.docx"
doc.save(out)
print("Saved:", out)
